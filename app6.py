import streamlit as st

# 将这行移到所有 st 命令之前
st.set_page_config(page_title="复新Vis-数据新闻多智能体工作流", layout="wide")

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from openai import OpenAI
import os
from dotenv import load_dotenv
import json
import logging
from typing import Dict, Tuple, Optional, List, Any
import functools
import time
import concurrent.futures
from dataclasses import dataclass
import re  # Import the regular expression module
import requests
import asyncio
from crawl4ai import AsyncWebCrawler
from bs4 import BeautifulSoup
from pyecharts import options as opts
from pyecharts.charts import Bar, Pie, Line, Scatter, HeatMap, Tree, Sunburst
from pyecharts.charts import TreeMap, Boxplot
from pyecharts.globals import ThemeType
from streamlit_echarts import st_pyecharts
import numpy as np  # 确保导入numpy用于直方图计算
from io import BytesIO, StringIO
from docx import Document
from docx.shared import Inches
import base64

def search_with_tavily(query):
    """使用 Tavily API 搜索相关信息"""
    try:
        TAVILY_API_KEY = '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'
        TAVILY_API_URL = 'https://api.tavily.com/search'
        
        data = {
            "api_key": TAVILY_API_KEY,
            "query": query,
            "search_depth": "basic",
            "max_results": 5,
            "language": "zh"
        }
        
        response = requests.post(TAVILY_API_URL, json=data)
        response.raise_for_status()
        result = response.json()
        
        # 提取最多3个搜索结果
        if 'results' in result:
            return [
                {
                    'title': item.get('title', ''),
                    'content': item.get('content', '')[:500],
                    'url': item.get('url', '')
                }
                for item in result['results'][:3]
            ]
        return []
    except Exception as e:
        logger.error(f"Tavily API 错误: {str(e)}")
        return []

# 新增数据收集智能体类
class DataCollectionAgent:
    """数据收集智能体，负责多维度数据收集和整合"""
    
    def __init__(self, client):
        self.client = client
        self.tavily_api_key = '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'
        self.tavily_url = 'https://api.tavily.com/search'
    
    def collect_multi_dimensional_data(self, directions: List[str], topic: str) -> Dict[str, Any]:
        """多维度数据收集"""
        results = {
            'structured_data': [],
            'text_data': [],
            'collection_summary': {},
            'failed_directions': []
        }
        
        progress = st.progress(0)
        status_text = st.empty()
        
        for i, direction in enumerate(directions):
            status_text.text(f"正在收集: {direction} ({i+1}/{len(directions)})")
            
            try:
                # 使用Tavily搜索相关信息
                search_results = self._search_with_tavily_enhanced(direction, topic)
                
                if search_results:
                    # 对每个搜索结果进行智能分析和结构化
                    for result in search_results[:3]:  # 限制每个方向最多处理3个结果
                        structured_data = self._intelligent_structurize(result, direction, topic)
                        
                        if structured_data is not None and not structured_data.empty:
                            structured_data['data_direction'] = direction
                            structured_data['source_url'] = result.get('url', '')
                            results['structured_data'].append(structured_data)
                        else:
                            # 如果无法结构化，保存为文本数据
                            results['text_data'].append({
                                'direction': direction,
                                'content': result.get('content', ''),
                                'url': result.get('url', ''),
                                'title': result.get('title', '')
                            })
                    
                    results['collection_summary'][direction] = f"成功收集到 {len(search_results)} 条相关信息"
                else:
                    results['failed_directions'].append(direction)
                    results['collection_summary'][direction] = "未找到相关数据"
                    
            except Exception as e:
                results['failed_directions'].append(direction)
                results['collection_summary'][direction] = f"收集失败: {str(e)}"
            
            progress.progress((i + 1) / len(directions))
        
        status_text.text("数据收集完成")
        progress.empty()
        status_text.empty()
        
        return results
    
    def _search_with_tavily_enhanced(self, query: str, topic: str) -> List[Dict]:
        """增强版Tavily搜索"""
        try:
            # 结合选题和具体方向进行搜索，并尝试限定常见数据文件类型
            enhanced_query = (
                f"{topic} {query} 数据 统计 报告 filetype:csv OR filetype:xls OR filetype:xlsx OR filetype:pdf"
            )

            data = {
                "api_key": self.tavily_api_key,
                "query": enhanced_query,
                "search_depth": "advanced",  # 使用高级搜索
                "max_results": 8,
                "language": "zh",
                "include_domains": ["gov.cn", "stats.gov.cn", "xinhuanet.com", "people.com.cn"],  # 优先权威来源
                "exclude_domains": ["baidu.com", "so.com"]  # 排除搜索引擎页面
            }
            
            response = requests.post(self.tavily_url, json=data, timeout=30)
            response.raise_for_status()
            result = response.json()
            
            if 'results' in result:
                results = result['results']
                # 简单根据标题匹配选题或方向关键词以提升相关度
                filtered = [
                    r for r in results
                    if topic.lower() in r.get('title', '').lower()
                    or query.lower() in r.get('title', '').lower()
                ]
                return filtered if filtered else results
            return []
            
        except Exception as e:
            logger.error(f"Enhanced Tavily search error: {str(e)}")
            return []
    
    def _intelligent_structurize(self, search_result: Dict, direction: str, topic: str) -> Optional[pd.DataFrame]:
        """智能结构化处理 - 使用两阶段推理：GLM-4-PLUS筛选 + GLM-Z1-AIRX深度推理"""
        try:
            content = search_result.get('content', '')
            title = search_result.get('title', '')
            url = search_result.get('url', '')
            
            # 如果内容太短，尝试爬取完整页面
            if len(content) < 200:
                try:
                    response = requests.get(url, timeout=15)
                    soup = BeautifulSoup(response.text, 'html.parser')
                    content = soup.get_text()[:3000]  # 限制长度
                except:
                    pass
            
            # 第一阶段：使用GLM-4-PLUS筛选关键数据段落
            filter_system_prompt = f"""
            你是数据筛选专家，擅长从新闻文本中识别包含数值、统计信息、比例、趋势等关键数据的段落。
            
            任务：从文本中提取与"{topic}"相关的"{direction}"方面的关键数据段落。
            
            要求：
            1. 识别包含具体数字、百分比、统计数据的段落
            2. 保留包含时间、地区、分类等维度信息的段落
            3. 过滤掉纯描述性、无数据价值的内容
            4. 保持原文表述，不要修改数据
            
            输出格式：直接输出筛选后的关键段落，如果没有找到有价值的数据段落，输出"NO_KEY_DATA"。
            """
            
            filter_messages = [
                {"role": "system", "content": filter_system_prompt},
                {"role": "user", "content": f"标题：{title}\n内容：{content[:2500]}"}
            ]
            
            filter_response = self.client.chat_completions_create(filter_messages, temperature=0.2)
            
            if 'choices' not in filter_response or len(filter_response['choices']) == 0:
                return None
                
            key_data_segments = filter_response['choices'][0]['message']['content'].strip()
            
            if key_data_segments == "NO_KEY_DATA" or "NO_KEY_DATA" in key_data_segments:
                return None
            
            # 第二阶段：使用GLM-Z1-AIRX进行深度推理和结构化
            reasoning_system_prompt = f"""
            你是一个专业的数据分析推理专家。请深度分析以下关键数据段落，运用逻辑推理将文本信息转换为结构化的表格数据。
            
            推理任务：
            1. 分析数据的内在逻辑和关联关系
            2. 识别数据的层次结构和分类体系
            3. 推断隐含的数据关系和计算逻辑
            4. 设计最优的表格结构（列名、数据类型）
            5. 确保数据的完整性、一致性和准确性
            
            主题：{topic}
            数据方向：{direction}
            
            输出要求：
            - 设计合理的表格结构，列名使用中文
            - 每行代表一个完整的数据记录
            - 直接输出CSV格式，第一行为列名
            - 如果经过推理仍无法构建有效表格，输出"REASONING_FAILED"
            """
            
            reasoning_messages = [
                {"role": "system", "content": reasoning_system_prompt},
                {"role": "user", "content": f"关键数据段落：\n{key_data_segments}\n\n请运用推理分析，设计表格结构并提取数据。"}
            ]
            
            # 使用推理模型进行深度分析
            try:
                reasoning_response = self.client.chat_completions_create(
                    reasoning_messages, 
                    model="glm-z1-airx",  # 使用推理模型
                    temperature=0.1
                )
            except Exception as e:
                # 如果推理模型不可用，回退到GLM-4-PLUS
                logger.warning(f"GLM-Z1-AIRX不可用，回退到GLM-4-PLUS: {str(e)}")
                reasoning_response = self.client.chat_completions_create(
                    reasoning_messages, 
                    model="glm-4-plus",
                    temperature=0.1
                )
            
            if 'choices' in reasoning_response and len(reasoning_response['choices']) > 0:
                csv_content = reasoning_response['choices'][0]['message']['content'].strip()
                
                if csv_content == "REASONING_FAILED" or "REASONING_FAILED" in csv_content:
                    return None
                
                try:
                    # 清理CSV内容
                    csv_content = csv_content.replace('```csv', '').replace('```', '').strip()
                    df = pd.read_csv(StringIO(csv_content))
                    
                    # 验证数据框
                    if len(df) > 0 and len(df.columns) > 1:
                        # 添加元数据列
                        df['数据来源'] = title
                        df['来源URL'] = url
                        df['收集时间'] = pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')
                        return df
                    
                except Exception as e:
                    logger.error(f"CSV parsing error: {str(e)}")
                    return None
            
            return None
            
        except Exception as e:
            logger.error(f"Intelligent structurization error: {str(e)}")
            return None

def smart_merge_dataframes(dataframes: List[pd.DataFrame]) -> pd.DataFrame:
    """智能合并不同结构的数据框"""
    if not dataframes:
        return pd.DataFrame()
    
    if len(dataframes) == 1:
        return dataframes[0]
    
    try:
        # 尝试直接合并（如果列结构相似）
        return pd.concat(dataframes, ignore_index=True, sort=False)
    except Exception:
        # 如果直接合并失败，使用智能合并策略
        merged_data = []
        
        for df in dataframes:
            # 将每个数据框转换为标准格式
            for _, row in df.iterrows():
                record = {
                    '数据项': '',
                    '数值': '',
                    '单位': '',
                    '时间': '',
                    '分类': '',
                    '数据来源': row.get('数据来源', ''),
                    '来源URL': row.get('来源URL', ''),
                    'data_direction': row.get('data_direction', ''),
                    '收集时间': row.get('收集时间', '')
                }
                
                # 尝试从行数据中提取标准字段
                for col, val in row.items():
                    if col not in ['数据来源', '来源URL', 'data_direction', '收集时间']:
                        if pd.api.types.is_numeric_dtype(type(val)):
                            record['数值'] = val
                            record['数据项'] = col
                        else:
                            if not record['数据项']:
                                record['数据项'] = col
                            if '年' in str(val) or '月' in str(val) or '日' in str(val):
                                record['时间'] = val
                            else:
                                record['分类'] = val
                
                merged_data.append(record)
        
        return pd.DataFrame(merged_data)

# 选题确定阶段
def topic_selection_phase():
    st.header("第一步：数据新闻选题确定")
    
    # 初始化session state变量
    if 'topic_conversation' not in st.session_state:
        st.session_state.topic_conversation = []
    
    if 'suggested_topics' not in st.session_state:
        st.session_state.suggested_topics = []
    
    if 'selected_topic' not in st.session_state:
        st.session_state.selected_topic = None
    
    if 'skip_topic_selection' not in st.session_state:
        st.session_state.skip_topic_selection = False
        
    # 如果已经选择了选题，显示它并进入下一阶段
    if st.session_state.selected_topic:
        st.success(f"已选择的选题：{st.session_state.selected_topic}")
        return True
    
    # 显示聊天历史
    for message in st.session_state.topic_conversation:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # 用户输入选题方向
    topic_description = st.chat_input(
        "请描述你感兴趣的数据新闻选题方向...",
        key="topic_input"
    )
    
    # 当用户提交选题描述
    if topic_description:
        # 添加用户消息到聊天历史
        st.session_state.topic_conversation.append({"role": "user", "content": topic_description})
        
        # 显示用户消息
        with st.chat_message("user"):
            st.markdown(topic_description)
        
        # 先尝试获取相关搜索结果
        search_results = search_with_tavily(topic_description)
        search_context = ""
        
        if search_results:
            search_context = "基于以下最新资讯:\n" + "\n\n".join([
                f"标题: {result['title']}\n内容: {result['content']}\n来源: {result['url']}"
                for result in search_results
            ])
        
        # 构建提示
        system_prompt = """你是一位专业有趣的数据新闻编辑，擅长帮助记者确定有价值的数据新闻选题。
        请根据用户的选题方向，生成三个明确具体的数据新闻选题建议。每个选题必须:
        1. 具有新闻价值和数据驱动特性
        2. 明确定义了研究问题和可能的数据来源
        3. 有潜在的社会影响或公众关注度
        
        按以下格式输出三个选题：
        [选题1]
        标题：(选题标题)
        核心问题：(选题要解决的核心问题)
        数据新闻价值：(为什么这个选题值得做数据新闻)
        
        [选题2]
        ...
        
        [选题3]
        ...
        """
        
        messages = [
            {"role": "system", "content": system_prompt}
        ]
        
        # 添加搜索上下文（如果有）
        if search_context:
            messages.append({"role": "system", "content": search_context})
        
        messages.append({"role": "user", "content": f"我想做一个关于以下主题的数据新闻：{topic_description}"})
        
        # 显示助手正在输入的消息
        with st.chat_message("assistant"):
            suggestion_text_container = st.empty()
            suggestion_text = ""
            
            # 使用智谱API的流式输出
            for token in client.chat_completions_create(messages, stream=True):
                suggestion_text += token
                suggestion_text_container.markdown(suggestion_text)
            
            # 保存完整回复到会话状态
            st.session_state.topic_conversation.append({"role": "assistant", "content": suggestion_text})
            
            # 解析建议的选题
            topics = []
            pattern = r'\[选题(\d+)\](.*?)(?=\[选题\d+\]|$)'
            matches = re.findall(pattern, suggestion_text, re.DOTALL)
            
            for _, topic_content in matches:
                # 提取选题信息
                title_match = re.search(r'标题：(.*?)(?:\n|$)', topic_content)
                title = title_match.group(1).strip() if title_match else "未命名选题"
                topics.append(title)
            
            st.session_state.suggested_topics = topics
    
    # 如果有建议的选题，提供选择按钮
    if st.session_state.suggested_topics:
        st.subheader("请选择一个选题，或重新生成")
        
        cols = st.columns(3)
        for i, topic in enumerate(st.session_state.suggested_topics):
            with cols[i]:
                if st.button(f"选择: {topic}"):
                    st.session_state.selected_topic = topic
                    st.rerun()
        
        if st.button("重新生成选题"):
            # 清除之前的建议，保留对话历史
            st.session_state.suggested_topics = []
            st.rerun()
    
    # 如果用户还没有选择选题，返回False
    return False

# 数据收集方向生成阶段
def data_collection_phase():
    st.header("第二步，整理数据收集思路")
    
    # 初始化session state变量
    if 'data_directions' not in st.session_state:
        st.session_state.data_directions = None
    
    if 'data_collection_completed' not in st.session_state:
        st.session_state.data_collection_completed = False
        
    if 'data_conversation' not in st.session_state:
        st.session_state.data_conversation = []
    
    # 如果已经完成数据收集方向生成，显示结果并隐藏输入框
    if st.session_state.data_collection_completed:
        for message in st.session_state.data_conversation:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        st.subheader("选择要自动收集的数据方向")

        parsed = parse_data_directions(st.session_state.data_directions)
        second_hand = parsed.get("二手数据", [])
        research = parsed.get("调研数据", []) + parsed.get("自主数据挖掘", [])

        topic = st.session_state.selected_topic

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("#### 二手数据检索")
            queries = [f"{topic} {d.strip()}" for d in second_hand]
            directions_input = st.text_area(
                "可编辑的检索问题（一行一个）",
                value="\n".join(queries),
                key="second_hand_input",
                help="这些检索问题将用于网络搜索"
            )
            if st.button("🚀 启动智能数据收集"):
                q_list = [d.strip() for d in directions_input.splitlines() if d.strip()]
                if q_list:
                    df = collect_data_from_directions(q_list)
                    if not df.empty:
                        processor = DataProcessor(df)
                        st.session_state['current_processor'] = processor
                        st.session_state['data_uploaded'] = True
                        st.success("✅ 智能数据收集完成并载入成功！")
                    else:
                        st.warning("⚠️ 未能获取到足够的结构化数据，请尝试调整数据收集方向或手动上传数据")

        with col2:
            if research:
                st.markdown("#### 调研/自主数据挖掘")
                for d in research:
                    st.write(f"- {d}")
                if st.button("生成问卷", key="gen_q"):
                    st.session_state.questionnaire = generate_questionnaire(research)
                if st.button("生成爬虫代码", key="gen_crawler"):
                    st.session_state.crawler_code = generate_crawler_code(research)
                if st.session_state.get('questionnaire'):
                    st.subheader("问卷示例")
                    st.markdown(st.session_state.questionnaire)
                if st.session_state.get('crawler_code'):
                    st.subheader("爬虫代码示例")
                    st.code(st.session_state.crawler_code, language='python')

        refresh_col1, refresh_col2 = st.columns([1, 10])
        with refresh_col1:
            if st.button("🔄", help="重新生成数据收集方向"):
                st.session_state.data_directions = None
                st.session_state.data_collection_completed = False
                st.rerun()
        with refresh_col2:
            st.write("如需重新生成数据收集方向，请点击左侧刷新按钮")

        return True
    
    # 当用户已经选择了选题，但还没有生成数据收集方向
    if st.session_state.selected_topic and not st.session_state.data_directions:
        topic = st.session_state.selected_topic
        
        # 显示已有对话
        for message in st.session_state.data_conversation:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # 系统提示：告知用户当前选题
        if len(st.session_state.data_conversation) == 0:
            with st.chat_message("assistant"):
                st.markdown(f"基于您选择的选题: **{topic}**，我可以帮您生成详细的数据收集方向。")
                st.session_state.data_conversation.append({
                    "role": "assistant", 
                    "content": f"基于您选择的选题: **{topic}**，我可以帮您生成详细的数据收集方向。"
                })
        
        # 用户输入或生成按钮
        user_input = st.chat_input("输入任何问题或点击'生成数据收集方向'按钮", key="data_input")
        generate_button = st.button("生成数据收集方向")
        
        if generate_button or user_input:
            if user_input:
                # 添加用户输入到对话
                st.session_state.data_conversation.append({"role": "user", "content": user_input})
                with st.chat_message("user"):
                    st.markdown(user_input)
                    
                # 进行普通回复
                with st.chat_message("assistant"):
                    response_container = st.empty()
                    response_text = ""
                    
                    # 构建普通对话提示
                    chat_messages = [
                        {"role": "system", "content": f"你是数据新闻专家，正在帮助用户规划选题'{topic}'的数据收集。回答用户所有关于数据收集的问题。"},
                    ]
                    
                    # 添加历史对话
                    for msg in st.session_state.data_conversation:
                        chat_messages.append({"role": msg["role"], "content": msg["content"]})
                    
                    # 使用流式API
                    for token in client.chat_completions_create(chat_messages, stream=True):
                        response_text += token
                        response_container.markdown(response_text)
                    
                    st.session_state.data_conversation.append({"role": "assistant", "content": response_text})
            else:
                # 生成数据收集方向
                # 构建提示
                system_prompt = """你是一位专业的数据新闻记者，擅长规划数据新闻报道的数据收集策略。

                首先，请判断用户选择的选题属于：
                - 📊 数据驱动型：从数据集出发，没有预设结论，通过数据探索发现故事
                - 💡 话题驱动型：基于明确的议题，收集数据来佐证或分析特定现象

                然后，根据用户选择的数据新闻选题，生成6-8个具体的数据检索方向设问，按照合理的新闻故事递进顺序排列：

                对每个数据收集方向，请注明：
                - 类型：1. 🌐 二手数据：提供可能存在相关数据的报告、具体网站、数据库或开放数据平台，附上数据获取方法。2. 🔍 调研数据：明确是需要线下走访、网络内容分析还是问卷发放，并提供调研的重点问题和方法。3. 🤖 自主数据挖掘：推荐适合爬虫收集的网站，说明可以获取什么类型的数据，以及大致的技术难度。
                - 该数据将回答什么具体问题
                - 数据获取的可行性评估（易/中/难）
                - 获取此数据可能遇到的挑战
                - 数据处理建议

                以Markdown格式输出，每个类别使用三级标题，每个具体方向使用四级标题，并使用表格或列表呈现详细信息。
                """
                
                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"请为我的数据新闻选题《{topic}》提供数据收集方向建议。"}
                ]
                
                # 显示进度条
                with st.chat_message("assistant"):
                    directions_container = st.empty()
                    directions_text = ""
                    
                    # 使用流式输出
                    for token in client.chat_completions_create(messages, stream=True):
                        directions_text += token
                        directions_container.markdown(directions_text)
                    
                    # 确保在流式输出完成后设置状态
                    st.session_state.data_directions = directions_text
                    st.session_state.data_collection_completed = True
                    st.session_state.data_conversation.append({
                        "role": "assistant", 
                        "content": directions_text
                    })
                
                # 强制重新加载页面以应用新状态
                st.rerun()
    
    return st.session_state.data_collection_completed

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
ZHIPU_API_KEY = '3a1df8f109f445f4b4eb898939a28a9f.0O5igS77SZZ0WGzV'  # 替换为您的API密钥
API_URL = 'https://open.bigmodel.cn/api/paas/v4/chat/completions'  # 智谱AI的API地址

# 修改OpenAI客户端初始化部分
class ZhipuClient:
    def __init__(self, api_key):
        self.api_key = api_key
        self.headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        }
        self.api_url = "https://open.bigmodel.cn/api/paas/v4/chat/completions"  # 智谱API URL

    def chat_completions_create(self, messages, model="glm-4-plus", temperature=0.7, stream=False):
        data = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
            "stream": stream
        }

        try:
            if not stream:
                # 非流式响应处理
                response = requests.post(self.api_url, headers=self.headers, json=data)
                response.raise_for_status()
                return response.json()
            else:
                # 流式响应处理
                response = requests.post(self.api_url, headers=self.headers, json=data, stream=True)
                response.raise_for_status()
                
                for line in response.iter_lines():
                    if line:
                        line = line.decode('utf-8')
                        if line.startswith('data: '):
                            json_str = line[6:]  # 去掉'data: '前缀
                            if json_str.strip() == '[DONE]':
                                break
                            try:
                                json_data = json.loads(json_str)
                                if 'choices' in json_data and len(json_data['choices']) > 0:
                                    content = json_data['choices'][0].get('delta', {}).get('content', '')
                                    if content:
                                        yield content
                            except json.JSONDecodeError:
                                pass
                        
        except Exception as e:
            if stream:
                yield f"API调用失败: {str(e)}"
            else:
                raise Exception(f"API调用失败: {str(e)}")

# 替换OpenAI客户端初始化
client = ZhipuClient(api_key=ZHIPU_API_KEY)

def get_data_summary(df: pd.DataFrame) -> str:
    """生成数据集的简要描述"""
    summary = []

    # 基本信息
    summary.append(f"数据集包含 {len(df)} 行，{len(df.columns)} 列")

    # 列信息
    for col in df.columns:
        col_type = df[col].dtype
        unique_count = df[col].nunique()
        null_count = df[col].isnull().sum()

        # 对于数值列，添加基本统计信息
        if pd.api.types.is_numeric_dtype(df[col]):
            stats = df[col].describe()
            col_info = (f"列 '{col}' (类型: {col_type}): "
                       f"取值范围 {stats['min']:.2f} 到 {stats['max']:.2f}, "
                       f"平均值 {stats['mean']:.2f}, "
                       f"不同值数量 {unique_count}")
        else:
            # 对于非数值列，显示唯一值数量和示例值
            sample_values = df[col].dropna().sample(min(3, unique_count)).tolist()
            col_info = (f"列 '{col}' (类型: {col_type}): "
                       f"不同值数量 {unique_count}, "
                       f"示例值: {', '.join(map(str, sample_values))}")

        if null_count > 0:
            col_info += f", 存在 {null_count} 个空值"

        summary.append(col_info)

    return "\n".join(summary)

def format_visualization_suggestions(response_text: str) -> str:
    """将API响应格式化为HTML样式的输出"""

    # 定义CSS样式
    css = """
        <style>
            .suggestion {
                background-color: #f8f9fa;
                padding: 15px 20px;
                margin-bottom: 20px;
                border-left: 4px solid #4A90E2;
                border-radius: 4px;
            }
            .suggestion-number {
                font-size: 18px;
                font-weight: bold;
                color: #2c3e50;
                margin-bottom: 15px;
            }
            .label {
                color: #4A90E2;
                font-weight: bold;
                margin-top: 10px;
            }
            .content {
                color: #333;
                margin-bottom: 10px;
                line-height: 1.5;
            }
        </style>
    """

    # 将文本分割成不同的建议
    suggestions = response_text.split("\n\n---\n\n")

    html_parts = [css]

    for i, suggestion in enumerate(suggestions, 1):
        # 开始新的建议区块
        html_parts.append(f'<div class="suggestion">')

        # 解析每个部分
        sections = suggestion.strip().split("\n\n")

        # 先添加建议编号
        html_parts.append(f'<div class="suggestion-number">建议 {i}</div>')

        # 处理每个部分
        for section in sections:
            if "[" in section and "]" in section:
                header = section[section.find("[")+1:section.find("]")]
                content = section[section.find("]")+1:].strip()
                html_parts.append(f'<div class="label">{header}</div>')
                html_parts.append(f'<div class="content">{content}</div>')

        html_parts.append('</div>')

    return "".join(html_parts)

# 修改get_llm_response函数中的prompt部分
def get_llm_response(prompt: str, df: Optional[pd.DataFrame] = None) -> str:
    """获取LLM的可视化建议，使用流式输出"""
    try:
        # 如果提供了DataFrame，生成数据概要
        if df is not None:
            data_summary = get_data_summary(df)
            full_prompt = f"""请作为数据可视化专家分析以下数据集：

数据集概要：
{data_summary}

用户问题：
{prompt}"""
        else:
            full_prompt = prompt

        # 构建完整的消息数组
        messages = [
            {
                "role": "system",
                "content": """你是一个中国数据新闻专家。请分析数据并提供3-4个具体的数据可视化建议。

每个建议必须按照以下固定格式输出，确保每个部分都另起新行：

[标题]
(带有探索性与新闻价值的标题)

[使用列]
(明确指出使用哪些列)

[图表类型]
(推荐使用的图表类型，如折线图、柱状图、散点图等)

[缘由]
(解释为什么这个可视化方案有价值)

---

建议 2：
(按相同格式继续...)"""
            },
            {"role": "user", "content": full_prompt}
        ]

        # 使用流式输出
        visualization_text = ""
        for token in client.chat_completions_create(messages, model="glm-4-plus", stream=True):
            visualization_text += token
        
        # 格式化可视化建议
        return format_visualization_suggestions(visualization_text)
    except Exception as e:
        logger.error(f"LLM API 错误: {str(e)}")
        return None

# 修改cached_api_call函数
@functools.lru_cache(maxsize=32)
def cached_api_call(prompt: str) -> str:
    """缓存API调用结果"""
    try:
        response = get_llm_response(prompt)
        if response is not None:
            return response
        else:
            st.error("无法获取AI建议，请稍后重试")
            return "无法获取AI建议，请稍后重试"
    except Exception as e:
        st.error(f"API调用错误: {str(e)}")
        logger.error(f"API调用错误: {str(e)}")
        return "API调用出现错误，请检查API密钥配置或网络连接"

# Data processing class
class DataProcessor:
    def __init__(self, file_or_df):
        self.df = None
        if isinstance(file_or_df, pd.DataFrame):
            # 直接使用DataFrame，与文件上传保持一致的处理方式
            self.df = file_or_df
            self.clean_data()  # 使用统一的数据清理方法
        else:
            self.file_type = file_or_df.name.split('.')[-1].lower()
            self.process_file(file_or_df)

    def clean_data(self):
        """统一的数据清理方法"""
        if self.df is not None:
            # 清理列名
            self.df.columns = self.df.columns.astype(str)
            self.df.columns = [col.strip() for col in self.df.columns]

            # 对每列进行基础处理
            for col in self.df.columns:
                # 处理日期时间列
                if any(keyword in col.lower() for keyword in ['time', 'date']):
                    try:
                        self.df[col] = pd.to_datetime(self.df[col])
                    except:
                        continue

                # 尝试转换为数值类型（如果适合的话）
                elif self.df[col].dtype == 'object':
                    try:
                        numeric_values = pd.to_numeric(self.df[col], errors='coerce')
                        if numeric_values.notna().sum() / len(numeric_values) > 0.5:
                            self.df[col] = numeric_values
                    except:
                        continue

    def process_file(self, file):
        """处理上传的文件"""
        try:
            # 读取文件
            if self.file_type == 'csv':
                self.df = pd.read_csv(file, encoding='utf-8')
            elif self.file_type == 'xlsx':
                self.df = pd.read_excel(file, engine='openpyxl')
            elif self.file_type == 'xls':
                self.df = pd.read_excel(file, engine='xlrd')
            elif self.file_type == 'json':
                self.df = pd.read_json(file)
            else:
                st.error("不支持的文件类型。请上传 CSV、XLSX、XLS 或 JSON 文件。")
                return

            self.clean_data()  # 使用统一的数据清理方法

        except UnicodeDecodeError:
            try:
                if self.file_type == 'csv':
                    self.df = pd.read_csv(file, encoding='gbk')
                    self.clean_data()
            except Exception as e:
                st.error(f"文件编码错误: {str(e)}")
                logger.error(f"文件编码错误: {str(e)}")
                raise
        except Exception as e:
            st.error(f"文件处理错误: {str(e)}")
            logger.error(f"文件处理错误: {str(e)}")
            raise

    def get_data_profile(self) -> Dict:
        """Generate basic data profile"""
        if self.df is None:  # Handle case where file processing failed
            return {}

        profile = {
            'columns': list(self.df.columns),
            'dtypes': {str(k): str(v) for k, v in self.df.dtypes.to_dict().items()},
            'null_counts': self.df.isnull().sum().to_dict(),
        }

        numeric_cols = self.df.select_dtypes(include=['int64', 'float64']).columns
        if not numeric_cols.empty:
            profile['statistics'] = {
                col: {
                    str(k): float(v) if pd.notnull(v) else None
                    for k, v in self.df[col].describe().to_dict().items()
                }
                for col in numeric_cols
            }

        return profile

# Visualization Generator (improved)
class VisualizationGenerator:
    def __init__(self, df: pd.DataFrame):
        self.df = df
        self.color_schemes = {
            'nyt': ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf'],
            'modern': ['#4A90E2', '#50E3C2', '#F5A623', '#D0021B', '#9013FE', '#417505', '#7ED321', '#BD10E0', '#8B572A', '#4A4A4A'],
            'soft': ['#ADD8E6', '#FF9999', '#FFB6C1', '#98D8C8', '#B0E0E6', '#FFDAB9', '#DDA0DD', '#E6E6FA', '#F0E68C', '#E0FFFF']
        }
        self.current_theme = 'modern'  # 默认主题
        self.theme_map = {
            'modern': ThemeType.LIGHT,
            'nyt': ThemeType.DARK,
            'soft': ThemeType.ESSOS
        }

    def set_theme(self, theme_name: str):
        """设置当前主题"""
        if theme_name in self.color_schemes:
            self.current_theme = theme_name

    def analyze_column(self, column: str) -> dict:
        series = self.df[column]
        unique_count = series.nunique()
        total_count = len(series)
        is_numeric = pd.api.types.is_numeric_dtype(series)

        return {
            'unique_count': unique_count,
            'total_count': total_count,
            'is_numeric': is_numeric,
            'dtype': str(series.dtype)
        }

    def preprocess_categorical_data(self, column: str) -> pd.DataFrame:
        """处理分类数据"""
        # 直接使用value_counts()获取分类统计
        value_counts = self.df[column].value_counts()
        
        # 如果类别过多，只保留前10个
        if len(value_counts) > 10:
            value_counts = value_counts.head(10)
        
        return pd.DataFrame({
            'category': value_counts.index,
            'count': value_counts.values
        })

    def suggest_chart_type(self, columns: List[str]) -> str:
        """根据数据特征自动推荐图表类型"""
        if len(columns) == 1:
            column = columns[0]
            analysis = self.analyze_column(column)
            
            if not analysis['is_numeric']:
                if analysis['unique_count'] <= 10:  # 少量分类
                    return 'pie' if analysis['unique_count'] <= 6 else 'bar'
                else:  # 大量分类
                    return 'bar'
            else:  # 数值数据
                if analysis['unique_count'] > 10:  # 连续数值
                    return 'histogram'
                else:  # 离散数值
                    return 'bar'
        else:  # 双变量分析
            x_col, y_col = columns[:2]
            x_analysis = self.analyze_column(x_col)
            y_analysis = self.analyze_column(y_col)
            
            if x_analysis['is_numeric'] and y_analysis['is_numeric']:
                return 'scatter'
            elif (not x_analysis['is_numeric']) and y_analysis['is_numeric']:
                return 'box' if x_analysis['unique_count'] <= 10 else 'violin'
            elif x_analysis['is_numeric'] and (not y_analysis['is_numeric']):
                return 'bar'
            else:
                return 'heatmap' if (x_analysis['unique_count'] <= 20 and y_analysis['unique_count'] <= 20) else 'bar'

    def generate_chart(self, columns: list, chart_type: str = 'auto', title: str = None, source: str = None, show_legend: bool = True):
        """使用pyecharts生成图表"""
        try:
            # 如果是自动模式，推荐图表类型
            if chart_type == 'auto':
             chart_type = self.suggest_chart_type(columns)

            # 设置默认宽高
            width = "100%"
            height = "500px"
        
            # 图表标题
            title_text = title or (f'{columns[0]} 分布' if len(columns) == 1 else f'{columns[1]} vs {columns[0]}')
        
            # 选择当前主题的颜色方案
            colors = self.color_schemes[self.current_theme]
            theme = self.theme_map[self.current_theme]
        
            # 设置图例选项
            legend_opts = opts.LegendOpts(is_show=show_legend, pos_bottom="10%", orient="horizontal", pos_left="center")
            
            if len(columns) == 1:
                column = columns[0]
                analysis = self.analyze_column(column)
                
                # 检查数据类型与图表类型的匹配性
                if not analysis['is_numeric'] and chart_type in ['histogram', 'box', 'violin', 'heatmap']:
                    st.warning("提示：文本/分类数据不适合使用数值型图表（直方图、箱线图、小提琴图、热力图）进行展示。请选择饼图、柱状图等分类图表。")
                    return None
                
                if not analysis['is_numeric']:
                    processed_data = self.preprocess_categorical_data(column)
                    
                    if chart_type == 'pie':
                        # 创建饼图
                        chart = (
                            Pie(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add(
                                series_name=column,
                                data_pair=[list(z) for z in zip(processed_data['category'], processed_data['count'])],
                                radius=["40%", "70%"],
                            )
                            .set_global_opts(
                                title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                                legend_opts=opts.LegendOpts(pos_bottom="10%", orient="horizontal", pos_left="center"),
                            )
                            .set_series_opts(label_opts=opts.LabelOpts(formatter="{b}: {c} ({d}%)"))
                        )
                        
                    elif chart_type == 'bar':
                        # 创建柱状图
                        chart = (
                            Bar(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add_xaxis(processed_data['category'].tolist())
                            .add_yaxis(column, processed_data['count'].tolist())
                            .set_global_opts(
                                title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                                xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=0)),
                                legend_opts=opts.LegendOpts(is_show=False),
                            )
                        )
                        
                    elif chart_type == 'treemap':
                        # 创建树图
                        data = [{"name": str(c), "value": int(v)} for c, v in zip(processed_data['category'], processed_data['count'])]
                        chart = (
                            TreeMap(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add(
                                series_name=column,
                                data=data,
                                visual_min=0,
                                visual_max=max(processed_data['count']),
                                label_opts=opts.LabelOpts(position="inside"),
                            )
                            .set_global_opts(title_opts=opts.TitleOpts(title=title_text))
                        )
                    
                    elif chart_type == 'sunburst':
                        # 创建旭日图
                        data = [{"name": str(c), "value": int(v)} for c, v in zip(processed_data['category'], processed_data['count'])]
                        chart = (
                            Sunburst(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add(
                                series_name=column,
                                data_pair=data,
                                radius=[0, "90%"],
                            )
                            .set_global_opts(title_opts=opts.TitleOpts(title=title_text))
                        )
                
                elif analysis['is_numeric']:
                    if chart_type == 'histogram':
                        # 处理直方图，使用Bar实现
                        # 生成直方图数据
                        hist, bin_edges = np.histogram(self.df[column].dropna(), bins='auto')
                        bin_labels = [f"{bin_edges[i]:.2f}-{bin_edges[i+1]:.2f}" for i in range(len(bin_edges)-1)]
                        
                        chart = (
                            Bar(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add_xaxis(bin_labels)
                            .add_yaxis("频率", hist.tolist())
                            .set_global_opts(
                                title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                                xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(rotate=45)),
                            )
                        )
                    
                    elif chart_type == 'box':
                        # 箱线图数据准备
                        data = self.df[column].dropna().tolist()
                        chart = (
                            Boxplot(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add_xaxis([column])
                            .add_yaxis("", self._prepare_boxplot_data(data))
                            .set_global_opts(
                                title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                                yaxis_opts=opts.AxisOpts(name=column),
                            )
                        )
                    
                    elif chart_type == 'violin':
                        # pyecharts不直接支持小提琴图，这里用boxplot替代
                        data = self.df[column].dropna().tolist()
                        st.warning("ECharts不直接支持小提琴图，已替换为箱线图展示")
                        chart = (
                            Boxplot(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                            .add_xaxis([column])
                            .add_yaxis("", self._prepare_boxplot_data(data))
                            .set_global_opts(
                                title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                                yaxis_opts=opts.AxisOpts(name=column),
                            )
                        )
                
            else:  # 双变量分析
                x_col, y_col = columns[:2]
                
                if chart_type == 'scatter':
                    # 散点图
                    chart = (
                        Scatter(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(self.df[x_col].tolist())
                        .add_yaxis(
                            y_col,
                            self.df[[x_col, y_col]].dropna().values.tolist(),
                            symbol_size=10,
                        )
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            xaxis_opts=opts.AxisOpts(name=x_col),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                            visualmap_opts=opts.VisualMapOpts(type_="size", max_=100, min_=10),
                        )
                    )
                
                elif chart_type == 'line':
                    # 折线图
                    chart = (
                        Line(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(self.df[x_col].tolist())
                        .add_yaxis(
                            y_col,
                            self.df[y_col].tolist(),
                            symbol_size=8,
                        )
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            xaxis_opts=opts.AxisOpts(
                                name=x_col,
                                type_="category" if not pd.api.types.is_numeric_dtype(self.df[x_col]) else "value"
                            ),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                            datazoom_opts=[opts.DataZoomOpts()],
                            legend_opts=opts.LegendOpts(pos_bottom="10%", orient="horizontal", pos_left="center"),
                        )
                    )
                
                elif chart_type == 'bar':
                    # 柱状图
                    chart = (
                        Bar(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(self.df[x_col].astype(str).tolist())
                        .add_yaxis(y_col, self.df[y_col].tolist())
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            xaxis_opts=opts.AxisOpts(
                                name=x_col,
                                axislabel_opts=opts.LabelOpts(rotate=45)
                            ),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                            datazoom_opts=[opts.DataZoomOpts()],
                            legend_opts=opts.LegendOpts(pos_bottom="10%", orient="horizontal", pos_left="center"),
                        )
                    )
                
                elif chart_type == 'box':
                    # 分组箱线图 - 简化实现
                    st.warning("ECharts中的分组箱线图实现较为复杂，展示效果可能与预期有差异")
                    grouped = self.df.groupby(x_col)[y_col].apply(list).reset_index()
                    chart = (
                        Boxplot(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(grouped[x_col].tolist())
                        .add_yaxis(
                            y_col,
                            [self._prepare_boxplot_data(data)[0] for data in grouped[y_col]]
                        )
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            xaxis_opts=opts.AxisOpts(name=x_col),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                        )
                    )
                
                elif chart_type == 'violin':
                    # 简化实现：pyecharts不直接支持小提琴图
                    st.warning("ECharts不直接支持小提琴图，已替换为分组箱线图展示")
                    grouped = self.df.groupby(x_col)[y_col].apply(list).reset_index()
                    chart = (
                        Boxplot(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(grouped[x_col].tolist())
                        .add_yaxis(
                            y_col,
                            [self._prepare_boxplot_data(data)[0] for data in grouped[y_col]]
                        )
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            xaxis_opts=opts.AxisOpts(name=x_col),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                        )
                    )
                
                elif chart_type == 'heatmap':
                    # 热力图
                    # 简化实现：创建数据透视表
                    if pd.api.types.is_numeric_dtype(self.df[y_col]):
                        # 如果y是数值列，计算平均值
                        pivot_data = self.df.pivot_table(
                            values=y_col,
                            index=x_col,
                            aggfunc='mean'
                        ).reset_index()
                        x_data = pivot_data[x_col].astype(str).tolist()
                        y_data = [y_col]
                        heat_data = [[0, 0, val] for val in pivot_data[y_col]]
                    else:
                        # 如果y是分类列，计算频数
                        counts = self.df.groupby([x_col, y_col]).size().reset_index(name='count')
                        x_data = sorted(counts[x_col].unique().astype(str).tolist())
                        y_data = sorted(counts[y_col].unique().astype(str).tolist())
                        heat_data = []
                        for _, row in counts.iterrows():
                            x_idx = x_data.index(str(row[x_col]))
                            y_idx = y_data.index(str(row[y_col]))
                            heat_data.append([x_idx, y_idx, row['count']])
                    
                    chart = (
                        HeatMap(init_opts=opts.InitOpts(width=width, height=height, theme=theme))
                        .add_xaxis(x_data)
                        .add_yaxis(
                            "",
                            y_data,
                            heat_data,
                        )
                        .set_global_opts(
                            title_opts=opts.TitleOpts(title=title_text, pos_left="center"),
                            visualmap_opts=opts.VisualMapOpts(),
                            xaxis_opts=opts.AxisOpts(name=x_col),
                            yaxis_opts=opts.AxisOpts(name=y_col),
                        )
                    )

            # 添加数据来源注释
            if source:
                chart.set_global_opts(
                    title_opts=opts.TitleOpts(
                        title=title_text,
                        subtitle=f"数据来源: {source}",
                        pos_left="center",
                        title_textstyle_opts=opts.TextStyleOpts(font_size=22)  # 设置更大的标题字号
                    )
                )

            return chart

        except Exception as e:
            st.error(f"生成图表时出错: {str(e)}")
            logger.error(f"图表生成错误: {str(e)}")
            return None

    # 辅助方法：为箱线图准备数据
    def _prepare_boxplot_data(self, data):
        """为ECharts箱线图准备数据"""
        result = []
        if not data:
            return [[0, 0, 0, 0, 0]]
            
        data = sorted(data)
        q1, q2, q3 = np.percentile(data, [25, 50, 75])
        iqr = q3 - q1
        low_whisker = max(min(data), q1 - 1.5 * iqr)
        high_whisker = min(max(data), q3 + 1.5 * iqr)
        result.append([low_whisker, q1, q2, q3, high_whisker])
        return result

    def evaluate_chart(self, chart_type: str, columns: List[str]) -> Tuple[str, List[str], str]:
        """评估图表适用性"""
        if chart_type == 'auto':
            return "非常适合", ["数据分析价值", "图表类型适用性"], "自动选择的图表类型最适合当前数据特征。"

        data_types = [self.df[col].dtype for col in columns]
        num_numeric = sum(1 for dtype in data_types if pd.api.types.is_numeric_dtype(dtype))
        num_categorical = sum(1 for dtype in data_types if dtype == 'object' or str(dtype).startswith('datetime'))
        num_columns = len(columns)
        score = "基本适合"  # 默认为"基本适合"
        feedback_dimensions = ["数据分析价值", "图表类型适用性"]
        feedback = ""

        if chart_type == 'line':
            # 检查X轴是否为时间类型
            if num_columns == 2:
                x_col = columns[0]
                is_time_col = pd.api.types.is_datetime64_any_dtype(self.df[x_col]) or \
                             any(keyword in x_col.lower() for keyword in ['time', 'date', '时间', '日期', '年', '月'])
                if not is_time_col:
                    score = "不适合"
                    feedback = "折线图最适合展示随时间变化的趋势。当前X轴不是时间类型，建议使用其他图表类型。"
                else:
                    score = "非常适合"
                    feedback = "折线图很好地展示了数据随时间的变化趋势。"
            else:
                score = "不适合"
                feedback = "折线图需要一个时间类型的X轴和一个数值类型的Y轴。"
        elif chart_type == 'pie':
            if num_columns != 1 or num_categorical != 1:
                score = "不适合"
                feedback = "饼图最适合展示单个分类变量的分布情况。"
            else:
                score = "非常适合"
                feedback = "饼图完美展现了单个分类变量的分布比例。"
        elif chart_type == 'bar':
            if num_columns == 1 and num_categorical == 1:
                score = "非常适合"
                feedback = "柱状图很好地展示了分类数据的对比。"
            elif num_columns == 2 and num_categorical == 1 and num_numeric == 1:
                score = "非常适合"
                feedback = "柱状图有效地展示了不同类别的数值对比。"
            else:
                score = "基本适合"
                feedback = "柱状图可以展示当前数据，但可能存在更好的可视化方式。"
        elif chart_type == 'scatter':
            if num_columns == 2 and num_numeric == 2:
                score = "非常适合"
                feedback = "散点图完美展示了两个数值变量间的关系。"
            else:
                score = "不适合"
                feedback = "散点图仅适用于展示两个数值变量的关系。"
        elif chart_type == 'histogram':
            if num_columns == 1 and num_numeric == 1:
                score = "非常适合"
                feedback = "直方图很好地展示了数值变量的分布情况。"
            else:
                score = "不适合"
                feedback = "直方图仅适用于展示单个数值变量的分布。"
        else:  # 其他图表类型
            score = "基本适合"
            feedback = "此图表类型可以展示当前数据。"

        return score, feedback_dimensions, feedback

    def get_chart_data(self, columns: List[str]) -> pd.DataFrame:
        """获取用于生成数据故事的图表数据"""
        if len(columns) == 1:
            column = columns[0]
            if not pd.api.types.is_numeric_dtype(self.df[column]):
                # 对于分类数据，计算频率和百分比
                value_counts = self.df[column].value_counts()
                percentages = value_counts / len(self.df) * 100
                return pd.DataFrame({
                    'category': value_counts.index,
                    'count': value_counts.values,
                    'percentage': percentages.values
                })
            else:
                # 对于数值数据，计算基本统计量
                stats = self.df[column].describe()
                return pd.DataFrame({
                    '统计指标': stats.index,
                    '值': stats.values
                })
        else:
            # 对于双变量分析，返回原始数据的相关部分
            return self.df[columns].copy()

def simulate_progress_bar():
    """模拟进度条动画"""
    progress_bar = st.progress(0)
    progress_text = st.empty()
    progress = 0

    while progress < 90:
        # 非线性进度增加，开始快，后面慢
        increment = max(0.3, (90 - progress) / 50)
        progress = min(90, progress + increment)

        # 更新进度条和文本
        progress_bar.progress(int(progress))
        progress_text.text(f'分析进度：{int(progress)}%')
        time.sleep(0.2)

    return progress_bar, progress_text

def get_data_story(chart_config: dict, data: pd.DataFrame, evaluation_score: str) -> str:
    """生成数据故事"""
    try:
        if data.empty:
            st.warning("没有足够的数据来生成故事。")
            return None

        # 构建数据概要字符串
        data_summary = "数据分析结果：\n"
        
        # 根据数据类型构建不同的描述
        if 'percentage' in data.columns:
            # 分类数据的描述
            total_count = data['count'].sum()
            data_summary += f"总计样本数：{total_count}\n\n"
            data_summary += "类别分布：\n"
            for _, row in data.iterrows():
                data_summary += f"- {row['category']}: {row['count']}次 (占比{row['percentage']}%)\n"
        
        elif '统计指标' in data.columns:
            # 数值数据的描述
            data_summary += "数值统计：\n"
            for _, row in data.iterrows():
                data_summary += f"- {row['统计指标']}: {row['值']}\n"
        
        else:
            # 其他数据类型的描述
            data_summary += data.to_string()

        # 构建Prompt
        prompt = f"""作为一名专业的数据新闻记者，请基于以下图表信息撰写一段数据新闻段落。
图表信息：
- 标题：{chart_config.get('title', '')}
- 图表类型：{chart_config.get('chart_type', '')}
- 使用数据列：{', '.join(chart_config.get('columns', []))}
- 数据来源：{chart_config.get('source', '')}

{data_summary}

要求：
1. 使用中文数据新闻专业写作风格
2. 突出数据发现的新闻价值
3. 客观陈述，准确引用数据
4. 注重数据背后的故事性
5. 语言简洁专业"""

        logger.info(f"发送给智谱AI的Prompt: \n{prompt}")

        # 调用智谱AI
        response = client.chat_completions_create(
            messages=[
                {"role": "system", "content": "你是一位经验丰富的数据新闻记者，擅长将数据分析转化为引人入胜的新闻故事。"},
                {"role": "user", "content": prompt}
            ],
            model="glm-4-plus",
            temperature=0.7
        )

        if 'choices' in response and len(response['choices']) > 0:
            story = response['choices'][0]['message']['content']
            logger.info(f"获得的故事内容: {story}")
            return story
        else:
            logger.error("API响应格式错误")
            st.error("生成故事时发生错误，API响应格式不正确。")
            return None

    except Exception as e:
        logger.error(f"生成故事时出错: {str(e)}")
        st.error(f"生成数据故事时发生错误: {str(e)}")
        return None

def get_data_news_story(selected_charts):
    """基于多个选定图表生成完整的数据新闻故事"""
    if not selected_charts:
        return None
    
    # 构建数据概要字符串
    charts_info = []
    
    for i, chart_info in enumerate(selected_charts):
        config = chart_info['config']
        data = chart_info['data']
        
        chart_summary = f"图表{i+1}信息：\n"
        chart_summary += f"- 标题：{config.get('title', f'图表{i+1}')}\n"
        chart_summary += f"- 图表类型：{config.get('chart_type', '自动')}\n"
        chart_summary += f"- 使用数据列：{', '.join(config.get('columns', []))}\n"
        chart_summary += f"- 数据来源：{config.get('source', '未指定')}\n\n"
        
        # 添加数据描述
        chart_summary += "数据分析结果：\n"
        
        if not data.empty:
            # 根据数据类型构建不同的描述
            if 'percentage' in data.columns:
                # 分类数据的描述
                total_count = data['count'].sum()
                chart_summary += f"总计样本数：{total_count}\n"
                chart_summary += "类别分布：\n"
                for _, row in data.head(5).iterrows():
                    chart_summary += f"- {row['category']}: {row['count']}次 (占比{row['percentage']:.2f}%)\n"
                if len(data) > 5:
                    chart_summary += f"- ... 等{len(data)}个类别\n"
            
            elif '统计指标' in data.columns:
                # 数值数据的描述
                chart_summary += "数值统计：\n"
                for _, row in data.iterrows():
                    chart_summary += f"- {row['统计指标']}: {row['值']}\n"
            
            else:
                # 其他数据类型的描述
                chart_summary += "数据预览：\n"
                chart_summary += data.head(3).to_string() + "\n...(数据省略)\n"
        
        charts_info.append(chart_summary)
    
    # 构建优化后的Prompt
    prompt = f"""作为一名专业数据新闻记者，请基于以下{len(charts_info)}个图表信息撰写一篇完整的数据新闻文章。

{"\n\n".join(charts_info)}

需求：
1. 文章需要一个吸引人的标题，使用"# 标题"格式
2. 将文章分成2-4个小节，每个小节标题使用"### 小节标题"格式
3. 文章篇幅适中（400-1100字），字数与图表数量成正比
4. 文章风格：
   - 开头引出核心发现，设置新闻基调
   - 中间部分深入分析每个图表数据，揭示数据背后的故事和关联
   - 结尾提供总结性观点或建议
5. 表达要求：
   - 客观准确引用数据，避免过度推测
   - 使用专业但通俗易懂的语言
   - 适当运用比喻、对比等修辞手法增强可读性

请直接输出完整的新闻文章，无需解释你的写作过程。"""

    logger.info(f"发送给智谱AI的Prompt: \n{prompt}")

    # 调用智谱AI
    try:
        response = client.chat_completions_create(
            messages=[
                {"role": "system", "content": "你是一位经验丰富的数据新闻记者，擅长将多维度数据分析转化为引人入胜的新闻故事。你会分析多个图表之间的关联，提炼出数据背后的深层含义。"},
                {"role": "user", "content": prompt}
            ],
            model="glm-4-plus",
            temperature=0.7
        )

        if 'choices' in response and len(response['choices']) > 0:
            story = response['choices'][0]['message']['content']
            logger.info(f"获得的新闻故事内容: {story}")
            return story
        else:
            logger.error("API响应格式错误")
            return None
    except Exception as e:
        logger.error(f"生成数据新闻时出错: {str(e)}")
        return None

# --- 新增功能：从文本中提取结构化数据 ---
def extract_structured_from_text(text: str) -> pd.DataFrame:
    """利用LLM将网页文本转换为结构化数据"""
    try:
        prompt = (
            "请从下面的新闻文本中提取与统计数字相关的数据，"
            "以CSV格式返回，第一行应为列名。\n\n" + text[:2000]
        )
        messages = [
            {"role": "system", "content": "你擅长从中文新闻文本中提取表格数据并以CSV形式输出"},
            {"role": "user", "content": prompt},
        ]
        response = client.chat_completions_create(messages)
        if 'choices' in response and len(response['choices']) > 0:
            csv_text = response['choices'][0]['message']['content']
            try:
                df = pd.read_csv(StringIO(csv_text))
                return df
            except Exception:
                logger.error("解析CSV失败")
                return pd.DataFrame()
    except Exception as e:
        logger.error(f"结构化解析失败: {e}")
    return pd.DataFrame()

# 解析模型生成的数据收集方向，按类别返回列表
def parse_data_directions(text: str) -> Dict[str, List[str]]:
    """将数据收集建议文本解析为类别 -> 方向列表的结构"""
    sections = re.findall(r"###\s*(.+?)\n(.*?)(?=\n###|$)", text, re.DOTALL)
    result: Dict[str, List[str]] = {}
    for name, content in sections:
        lines = re.findall(r"####\s*(.+)", content)
        if not lines:
            lines = re.findall(r"-\s*(.+)", content)
        if lines:
            result[name.strip()] = [l.strip() for l in lines]
    return result


def generate_questionnaire(directions: List[str]) -> str:
    """调用模型生成不超过15题的调研问卷"""
    prompt = (
        "请根据以下调研方向设计一份不超过15题的问卷，直接列出问题列表：\n" + "\n".join(directions)
    )
    messages = [
        {"role": "system", "content": "你是一名经验丰富的问卷设计专家"},
        {"role": "user", "content": prompt},
    ]
    resp = client.chat_completions_create(messages)
    if 'choices' in resp and resp['choices']:
        return resp['choices'][0]['message']['content']
    return "问卷生成失败"


def generate_crawler_code(directions: List[str]) -> str:
    """根据自主数据挖掘需求生成简单的爬虫示例代码"""
    prompt = (
        "请依据以下网站数据挖掘需求，提供一个Python爬虫示例，使用requests和BeautifulSoup，将结果保存为CSV：\n"
        + "\n".join(directions)
    )
    messages = [
        {"role": "system", "content": "你是一名擅长编写网络爬虫的Python开发者"},
        {"role": "user", "content": prompt},
    ]
    resp = client.chat_completions_create(messages)
    if 'choices' in resp and resp['choices']:
        return resp['choices'][0]['message']['content']
    return "爬虫代码生成失败"

# 根据多个数据方向自动收集网络数据并合并
def collect_data_from_directions(directions: List[str]) -> pd.DataFrame:
    """优化的数据收集流程"""
    if not st.session_state.get('selected_topic'):
        st.error("请先选择一个数据新闻选题")
        return pd.DataFrame()
    
    topic = st.session_state.selected_topic
    
    # 使用数据收集智能体
    agent = DataCollectionAgent(client)
    
    st.info("🤖 启动智能数据收集代理...")
    
    # 执行多维度数据收集
    collection_results = agent.collect_multi_dimensional_data(directions, topic)
    
    # 保存结构化数据列表到session state
    if collection_results['structured_data']:
        st.session_state.structured_data_list = collection_results['structured_data']
    
    # 显示收集摘要
    with st.expander("📊 数据收集摘要", expanded=True):
        for direction, summary in collection_results['collection_summary'].items():
            if direction in collection_results['failed_directions']:
                st.error(f"❌ {direction}: {summary}")
            else:
                st.success(f"✅ {direction}: {summary}")
    
    # 合并所有结构化数据
    all_structured_data = collection_results['structured_data']
    
    if all_structured_data:
        try:
            # 为每个数据表创建独立的展示
            st.subheader("📊 收集到的结构化数据")
            
            tabs = st.tabs([f"数据表 {i+1}" for i in range(len(all_structured_data))])
            
            for i, (tab, df) in enumerate(zip(tabs, all_structured_data)):
                with tab:
                    st.write(f"**数据方向**: {df.get('data_direction', ['未知'])[0] if not df.empty and 'data_direction' in df.columns else '未知'}")
                    st.write(f"**数据来源**: {df.get('数据来源', ['未知'])[0] if not df.empty and '数据来源' in df.columns else '未知'}")
                    display_df = df.drop(['data_direction'], axis=1, errors='ignore')
                    st.dataframe(display_df, use_container_width=True)
            
            # 提供多Sheet下载
            excel_data = export_multi_sheet_data(all_structured_data, topic)
            st.download_button(
                label="📥 下载多Sheet数据表",
                data=excel_data,
                file_name=f"{topic}_多维度数据.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            
            # 智能合并不同结构的数据框
            final_df = smart_merge_dataframes(all_structured_data)
            
            st.success(f"🎉 成功收集并结构化了 {len(all_structured_data)} 个数据集，合并后共 {len(final_df)} 行数据")
            
            return final_df
            
        except Exception as e:
            st.error(f"数据合并失败: {str(e)}")
            logger.error(f"Data merging error: {str(e)}")
    
    # 如果没有结构化数据，显示文本数据
    if collection_results['text_data']:
        st.warning("⚠️ 未能获取到结构化数据，但收集到了相关文本信息")
        
        with st.expander("📄 收集到的文本信息", expanded=False):
            for text_data in collection_results['text_data']:
                st.write(f"**方向**: {text_data['direction']}")
                st.write(f"**标题**: {text_data['title']}")
                st.write(f"**内容**: {text_data['content'][:300]}...")
                st.write(f"**来源**: {text_data['url']}")
                st.write("---")
    
    return pd.DataFrame()

# 新增爬虫数据处理类
class WebDataCrawler:
    def __init__(self):
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
    
    def crawl_data(self, url: str) -> pd.DataFrame:
        """爬取网页数据并转换为DataFrame"""
        try:
            st.write(f"开始爬取数据: {url}")
            progress_bar = st.progress(0)
            
            # 发送请求获取页面内容
            response = requests.get(url, headers=self.headers, timeout=30)
            response.raise_for_status()
            
            progress_bar.progress(0.5)
            
            # 尝试提取表格数据
            tables = pd.read_html(response.text)
            if tables:
                df = tables[0]  # 获取第一个表格
                # 保存到session state
                st.session_state['crawled_df'] = df
                st.write("数据预览：")
                st.write(df.head())
                progress_bar.progress(1.0)
                return df
            
            # 如果没有表格，尝试从文本中提取结构化数据
            soup = BeautifulSoup(response.text, 'html.parser')
            text_content = soup.get_text()
            df = extract_structured_from_text(text_content)
            if df is None or df.empty:
                df = pd.DataFrame({
                    'content': [text_content],
                    'url': [url],
                    'timestamp': [pd.Timestamp.now()]
                })
            st.session_state['crawled_df'] = df
            progress_bar.progress(1.0)
            st.write("获取到的文本数据预览：")
            st.write(text_content[:500] + "...")
            return df
            
        except Exception as e:
            msg = getattr(e, 'response', None)
            if msg is not None:
                err_detail = f"{msg.status_code} {msg.reason}"
            else:
                err_detail = str(e)
            st.error(f"爬取失败: {err_detail}")
            logger.error(f"爬取失败: {err_detail}")
            return pd.DataFrame()

# Main Application
def main():
    # 初始化session state变量
    if 'selected_charts' not in st.session_state:
        st.session_state.selected_charts = []
    
    st.title("复新Vis-数据新闻多智能体工作流")
    
    # 使用纯白色背景，只保留介绍文字的样式
    st.markdown(
        """
        <style>
        .intro-text {
            padding: 20px;
            border-radius: 10px;
            background-color: rgba(255, 255, 255, 0.9);
            margin: 20px 0;
        }
        
        .intro-point {
            margin: 10px 0;
            padding-left: 20px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # 始终显示侧边栏
    with st.sidebar:
        st.header("当你准备好了，你可以开始数据输入")
        data_input_method = st.radio(
            "选择数据输入方式",
            ["上传文件", "网页爬取"]
        )
        
        if data_input_method == "上传文件":
            uploaded_file = st.file_uploader("上传 CSV、Excel 或 JSON 文件", 
                                           type=['csv', 'xlsx', 'xls', 'json'])
            if uploaded_file:
                processor = DataProcessor(uploaded_file)
                # 保存到session state
                st.session_state['current_processor'] = processor
                # 设置状态，表示已上传数据
                st.session_state['data_uploaded'] = True
                
        else:  # 网页爬取
            url = st.text_input("输入要爬取的网页URL")
            
            with st.expander("爬取配置"):
                timeout = st.slider("超时时间(秒)", 10, 60, 30)
            
            if st.button("开始爬取", key="crawl_button"):
                if url:
                    try:
                        crawler = WebDataCrawler()
                        df = crawler.crawl_data(url)
                        
                        if not df.empty:
                            processor = DataProcessor(df)
                            # 保存到session state
                            st.session_state['current_processor'] = processor
                            # 设置状态，表示已上传数据
                            st.session_state['data_uploaded'] = True
                            st.success("数据爬取成功！")
                        else:
                            st.warning("未获取到数据，请检查URL或尝试其他网页。")
                    except Exception as e:
                        st.error(f"爬取失败: {str(e)}")

    # 从session state获取processor
    processor = st.session_state.get('current_processor', None)
    
    # 如果数据未上传且未跳过选题阶段，则显示选题和数据收集界面
    if not st.session_state.get('data_uploaded', False):
        # 第一阶段：选题确定
        if not topic_selection_phase():
            return  # 如果还没完成选题确定，不进入下一阶段
        
        # 第二阶段：数据收集方向
        if not data_collection_phase():
            return  # 如果还没完成数据收集方向生成，不进入下一阶段
            
        # 如果仍然没有数据，显示介绍内容
        if not processor or processor.df is None:
            st.markdown(
                """
                <div class="intro-text">
                    <h3>接下来的工作流程（版本0226）</h3>
                    <div class="intro-point">📊 <b>首先，上传你的数据：</b>支持上传本地数据集或使用我们的网页数据爬取</div>
                    <div class="intro-point">🤖 <b>然后，获取可视化的建议：</b>大模型会基于数据特征提供专业的可视化建议</div>
                    <div class="intro-point">📈 <b>其次，制作可视化图表：</b>提供多种图表类型和漂亮的配色选择</div>
                    <div class="intro-point">📝 <b>最后，拿上评估合格的图表，撰写出数据故事：</b>自动生成专业媒体风格的数据新闻段落</div>
                </div>
                """,
                unsafe_allow_html=True
            )
            return
    
    # 如果上传了数据，显示数据处理和可视化界面
    if processor and processor.df is not None:
        # 如果之前完成了选题阶段，显示选题信息
        if st.session_state.get('selected_topic') and not st.session_state.get('skip_topic_selection'):
            with st.expander("已选定的选题", expanded=False):
                st.success(f"数据新闻选题：{st.session_state.selected_topic}")
                if st.session_state.get('data_directions'):
                    st.markdown(st.session_state.data_directions)
        
        # 1. 数据预览部分 - 默认展开
        with st.expander("数据预览", expanded=True):
            st.dataframe(processor.df.head(31), use_container_width=True)

        # 2. 数据可视化建议部分
        st.subheader("第三步，获取数据可视化建议")
        suggestion_container = st.container()
        with suggestion_container:
            if st.button("获取可视化建议", key="viz_suggestion_btn"):
                progress_bar, progress_text = simulate_progress_bar()

                response = get_llm_response("请为这个数据集提供可视化建议", processor.df)

                # 完成时将进度设为100%
                progress_bar.progress(100)
                progress_text.text('分析完成！')
                time.sleep(0.5)  # 短暂显示完成状态

                progress_bar.empty()
                progress_text.empty()

                if response:
                    st.session_state.visualization_suggestions = response
                    st.markdown(response, unsafe_allow_html=True)

            elif st.session_state.get('visualization_suggestions'):
                st.markdown(st.session_state.visualization_suggestions, unsafe_allow_html=True)

        # 3. 可视化制作部分
        st.subheader("第四步，创建可视化")
        col1, col2 = st.columns([1, 2])

        with col1:
            # 添加主题选择
            color_theme = st.selectbox(
                "选择配色主题",
                options=['modern', 'nyt', 'soft'],
                format_func=lambda x: {
                    'modern': '现代简约',
                    'nyt': '新闻专业',
                    'soft': '柔和清新'
                }[x]
            )

            show_legend = st.checkbox("显示图例", value=True)

            viz_type = st.radio(
                "选择分析类型",
                options=['单列分析', '双列关系分析'],
                horizontal=True
            )

            custom_title = st.text_input("输入图表标题（可选）", "")
            data_source = st.text_input("输入数据来源（可选）", "")

            if viz_type == '单列分析':
                column = st.selectbox("选择要分析的列", options=processor.df.columns)
                chart_type = st.radio(
                    "选择图表类型",
                    options=['自动', '饼图', '柱状图', '直方图', '箱线图', '小提琴图', '树图', '旭日图'],
                    horizontal=True
                )
                columns_to_use = [column]
            else:
                x_column = st.selectbox("选择 X 轴数据", options=processor.df.columns)
                y_column = st.selectbox("选择 Y 轴数据", options=processor.df.columns)
                chart_type = st.radio(
                    "选择图表类型",
                    options=['自动', '折线图', '柱状图', '散点图', '箱线图', '小提琴图', '热力图'],
                    horizontal=True
                )
                columns_to_use = [x_column, y_column]

            if st.button("生成图表"):
                # 每次生成新图表时，清除之前的数据故事
                if 'data_story' in st.session_state:
                    del st.session_state['data_story']
                
                st.session_state.show_legend = show_legend

                st.session_state.current_chart_config = {
                    'viz_type': viz_type,
                    'columns': columns_to_use,
                    'chart_type': chart_type,
                    'title': custom_title,
                    'source': data_source
                }

        with col2:
            if 'current_chart_config' in st.session_state:
                config = st.session_state.current_chart_config
                vis_gen = VisualizationGenerator(processor.df)
                vis_gen.set_theme(color_theme)  # 设置选择的主题

                # 转换英文图表类型为中文
                chart_type_map = {
                    '自动': 'auto',
                    '饼图': 'pie',
                    '柱状图': 'bar',
                    '直方图': 'histogram',
                    '折线图': 'line',
                    '散点图': 'scatter',
                    '箱线图': 'box',
                    '小提琴图': 'violin',
                    '树图': 'treemap',
                    '旭日图': 'sunburst'
                }

                chart_type = chart_type_map.get(config['chart_type'], config['chart_type'])

                chart = vis_gen.generate_chart(
                    columns=config['columns'],
                    chart_type=chart_type,
                    title=config['title'],
                    source=config['source'],
                    show_legend=st.session_state.get('show_legend', True)  # 获取图例显示状态
                )

                if chart:
                    # 修改：使用st_pyecharts显示ECharts图表，而不是st.plotly_chart
                    st_pyecharts(chart, height="500px")

                    with st.expander("图表评估结果（对图表点击右键可保存为图片）", expanded=True):
                        score, dimensions, feedback = vis_gen.evaluate_chart(
                            chart_type,
                            config['columns']
                        )
                        st.write(f"**图表评估得分:** {score}")
                        st.write("**评估维度:**")
                        for dim in dimensions:
                            st.write(f"- {dim}")
                        st.write(f"**评估建议:** {feedback}")

                        # 移除原有的故事生成按钮，替换为选定按钮
                        if score in ["基本适合", "非常适合"]:
                            if st.button("选定此图表"):
                                # 检查是否已经选定了5个图表
                                if len(st.session_state.selected_charts) >= 5:
                                    st.warning("最多只能选定5个图表！请先删除一些已选定的图表。")
                                else:
                                    # 将当前图表配置和数据添加到已选定图表列表
                                    chart_data = vis_gen.get_chart_data(config['columns'])
                                    # 保存图表配置、数据和评估信息
                                    chart_info = {
                                        'config': config.copy(),
                                        'data': chart_data,
                                        'score': score,
                                        'chart': chart  # 保存图表对象
                                    }
                                    st.session_state.selected_charts.append(chart_info)
                                    st.success(f"已选定图表，当前已选定 {len(st.session_state.selected_charts)} 个图表")

    # 显示已选定的图表和第三步生成数据新闻
    if processor and processor.df is not None:  # 确保有数据被加载后才执行
        if 'selected_charts' in st.session_state and st.session_state.selected_charts:
            st.subheader("已选定的图表")
            # 使用列表容器展示已选图表
            for i, chart_info in enumerate(st.session_state.selected_charts):
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    st.write(f"**图表 {i+1}**: {chart_info['config'].get('title', '未命名图表')}")
                    # 显示图表
                    st_pyecharts(chart_info['chart'], height="300px")
                
                with col2:
                    st.write(f"评估: {chart_info['score']}")
                    # 添加删除按钮
                    if st.button(f"删除此图表", key=f"del_chart_{i}"):
                        st.session_state.selected_charts.pop(i)
                        st.rerun()
            
            # 第三步 - 生成完整数据新闻
            st.subheader("第五步，写作数据故事")
            
            if st.button("生成完整数据新闻"):
                progress_bar = st.progress(0)
                progress_text = st.empty()
                
                # 模拟进度
                for i in range(0, 101, 10):
                    progress_bar.progress(i)
                    if i == 0:
                        progress_text.text("初始化数据分析...")
                    elif i == 20:
                        progress_text.text("提取数据关键点...")
                    elif i == 40:
                        progress_text.text("构建新闻故事架构...")
                    elif i == 60:
                        progress_text.text("生成新闻内容...")
                    elif i == 80:
                        progress_text.text("润色文章表达...")
                    time.sleep(1.7)
                
                # 实际生成数据新闻
                story = get_data_news_story(st.session_state.selected_charts)
                
                # 完成进度
                progress_bar.progress(100)
                progress_text.text("数据新闻生成完成！")
                time.sleep(0.5)
                
                # 清除进度条和文本
                progress_bar.empty()
                progress_text.empty()
                
                if story:
                    st.session_state.news_story = story
                else:
                    st.error("无法生成数据新闻，请稍后重试。")
            
            # 显示数据新闻
            if 'news_story' in st.session_state:
                # 设计一个富媒体框来展示新闻内容
                st.markdown(
                    """
                    <style>
                    .news-container {
                        padding: 20px;
                        background-color: #f8f9fa;
                        border-radius: 10px;
                        border-left: 5px solid #4A90E2;
                        margin: 10px 0;
                    }
                    .news-title {
                        font-size: 24px;
                        font-weight: bold;
                        margin-bottom: 15px;
                        color: #2c3e50;
                    }
                    .news-section {
                        font-size: 18px;
                        font-weight: bold;
                        margin: 15px 0 10px 0;
                        color: #3498db;
                    }
                    .news-content {
                        font-size: 16px;
                        line-height: 1.6;
                        color: #333;
                    }
                    </style>
                    """,
                    unsafe_allow_html=True
                )
                
                # 处理Markdown格式的新闻内容
                news_content = st.session_state.news_story
                
                # 使用正则表达式提取标题和小节标题
                # 假设最大的标题使用# 或## 开始，小节标题使用### 开始
                title_match = re.search(r'^#\s+(.+)$|^##\s+(.+)$', news_content, re.MULTILINE)
                if title_match:
                    title = title_match.group(1) if title_match.group(1) else title_match.group(2)
                    # 从内容中移除主标题
                    news_content = re.sub(r'^#\s+.+$|^##\s+.+$', '', news_content, count=1, flags=re.MULTILINE)
                else:
                    title = "数据新闻报道"
                
                # 查找所有小节标题和内容
                sections = re.split(r'^###\s+(.+)$', news_content, flags=re.MULTILINE)
                
                # 显示富媒体格式的新闻
                news_html = f'<div class="news-container"><div class="news-title">{title}</div>'
                
                if len(sections) > 1:  # 有小节标题
                    for i in range(1, len(sections), 2):
                        if i < len(sections):
                            section_title = sections[i]
                            section_content = sections[i + 1] if i + 1 < len(sections) else ""
                            news_html += f'<div class="news-section">{section_title}</div>'
                            news_html += f'<div class="news-content">{section_content}</div>'
                else:  # 没有小节标题，直接显示内容
                    news_html += f'<div class="news-content">{news_content}</div>'
                
                news_html += '</div>'
                st.markdown(news_html, unsafe_allow_html=True)
                
                # 提供下载按钮
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="下载Markdown格式",
                        data=st.session_state.news_story,
                        file_name="data_news_story.md",
                        mime="text/markdown"
                    )
                with col2:
                    word_file = export_to_word(st.session_state.news_story)
                    st.download_button(
                        label="下载Word文档",
                        data=word_file,
                        file_name="data_news_story.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                # 添加数据导出功能
                if 'current_processor' in st.session_state and hasattr(st.session_state.current_processor, 'df'):
                    st.subheader("数据导出")
                    col3, col4 = st.columns(2)
                    
                    with col3:
                        # 导出当前数据集
                        csv_data = st.session_state.current_processor.df.to_csv(index=False)
                        st.download_button(
                            label="下载当前数据集(CSV)",
                            data=csv_data,
                            file_name="current_dataset.csv",
                            mime="text/csv"
                        )
                    
                    with col4:
                        # 如果有多个结构化数据集，提供多Sheet Excel下载
                        if hasattr(st.session_state, 'structured_data_list') and st.session_state.structured_data_list:
                            excel_file = export_multi_sheet_data(
                                st.session_state.structured_data_list, 
                                st.session_state.get('selected_topic', '数据分析')
                            )
                            st.download_button(
                                label="下载多Sheet数据(Excel)",
                                data=excel_file,
                                file_name="multi_sheet_data.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
        elif processor and processor.df is not None:  # 只在有数据但没有选定图表时显示提示
            st.info("请先选定至少一个图表，才能生成数据新闻。")

def export_multi_sheet_data(structured_data_list: List[pd.DataFrame], topic: str):
    """导出多Sheet Excel文件"""
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        sheet_names = set()
        
        for i, df in enumerate(structured_data_list):
            # 根据数据方向生成Sheet名称
            direction = df.get('data_direction', [f'数据{i+1}'])[0] if not df.empty else f'数据{i+1}'
            sheet_name = direction[:30]  # Excel sheet名称限制
            
            # 确保Sheet名称唯一
            original_name = sheet_name
            counter = 1
            while sheet_name in sheet_names:
                sheet_name = f"{original_name}_{counter}"
                counter += 1
            
            sheet_names.add(sheet_name)
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    
    output.seek(0)
    return output

def export_to_word(news_content, selected_charts=None):
    """将新闻内容导出为Word文档"""
    doc = Document()
    
    # 处理Markdown内容
    # 提取主标题
    title_match = re.search(r'^#\s+(.+)$|^##\s+(.+)$', news_content, re.MULTILINE)
    if title_match:
        title = title_match.group(1) if title_match.group(1) else title_match.group(2)
        doc.add_heading(title, level=0)
        # 从内容中移除主标题
        news_content = re.sub(r'^#\s+.+$|^##\s+.+$', '', news_content, count=1, flags=re.MULTILINE)
    
    # 处理小节标题和内容
    sections = re.split(r'^###\s+(.+)$', news_content, flags=re.MULTILINE)
    
    # 如果有小节
    if len(sections) > 1:
        # 处理第一个非标题部分(如果有的话)
        if sections[0].strip():
            doc.add_paragraph(sections[0].strip())
            
        # 处理各小节
        for i in range(1, len(sections), 2):
            if i < len(sections):
                section_title = sections[i]
                section_content = sections[i + 1] if i + 1 < len(sections) else ""
                
                # 添加小节标题
                doc.add_heading(section_title, level=2)
                
                # 添加小节内容
                doc.add_paragraph(section_content.strip())
    else:
        # 没有小节，直接添加内容
        doc.add_paragraph(news_content.strip())
    
    # 保存Word文档到内存
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"Application error: {str(e)}")
        logger.error(f"Application error: {str(e)}")
